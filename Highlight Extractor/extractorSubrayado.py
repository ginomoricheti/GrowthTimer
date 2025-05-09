
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def extract_highlighted_text(pdf_path):
    '''Extrae el texto resaltado junto con su color de fondo desde un archivo PDF.'''
    doc = fitz.open(pdf_path)
    highlighted_data = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        annot = page.first_annot

        while annot:
            if annot.type[0] == 8:  # Highlight
                rect = annot.rect
                text = page.get_textbox(rect).strip()

                # Detectar color del resaltado (por defecto amarillo)
                color_rgb = annot.colors.get('stroke', (1, 1, 0))
                color_hex = '#%02X%02X%02X' % tuple(int(c * 255) for c in color_rgb)

                if text:
                    highlighted_data.append((text, color_hex))

            annot = annot.next

    return highlighted_data

def set_run_background_color(run, hex_color):
    """Establece el color de fondo de un run en Word usando XML directo."""
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace('#', ''))
    rPr.append(shd)

# Guardar el texto resaltado en un archivo Word
def save_to_word(highlighted_data, output_path):
    '''Guarda los bloques resaltados en Word con una etiqueta textual según su color.'''
    doc = Document()

    color_map = {
        '#3E6FF5': '!nivel1',
        '#FF0000': '!nivel2',
        '#339E00': '!nivel3',
        '#FFED00': '!importante'
    }

    for text, color in highlighted_data:
        blocks = text.split('. ')
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            if not block.endswith('.'):
                block += '.'

            etiqueta = color_map.get(color.upper(), '!nivel1')
            block += f' {etiqueta}'

            para = doc.add_paragraph()
            para.add_run(block)

    doc.save(output_path)
    return output_path

# Reformatear el Word para eliminar espacios en blanco
def reformat_word_doc(input_path, output_path):
    '''🟡 #FFED00 - Reformatea el documento Word eliminando líneas vacías.'''
    document = Document(input_path)
    new_doc = Document()

    style = new_doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            new_doc.add_paragraph(text)

    new_doc.save(output_path)


# ================================
# 🟦 #3E6FF5 - TÍTULO NIVEL 1
# BUCLE PRINCIPAL DEL PROGRAMA
# ================================

def main():
    while True:
        print("""
==========================================
  EXTRACTOR DEL TEXTO REMARCADO EN EL PDF
        por Gino Morichetti
                >.<
==========================================
""")
        try:
            import_route = input('Ingrese la ruta del archivo PDF (o 0 para salir): ').strip()
            if import_route == '0':
                print('Programa finalizado.')
                break

            if not import_route.lower().endswith('.pdf'):
                print('El archivo debe tener extensión .pdf')
                continue

            if not os.path.isfile(import_route):
                print(f'El archivo "{import_route}" no se ha encontrado.')
                continue

            highlighted = extract_highlighted_text(import_route)

            if highlighted:
                # 🟢 #339E00 - Definir rutas de salida
                pdf_dir = os.path.dirname(import_route)
                base_name = os.path.splitext(os.path.basename(import_route))[0]

                temp_docx = os.path.join(pdf_dir, f'{base_name}_temporal.docx')
                final_docx = os.path.join(pdf_dir, f'{base_name}_resaltado.docx')

                save_to_word(highlighted, temp_docx)
                reformat_word_doc(temp_docx, final_docx)

                os.remove(temp_docx)
                print(f'\n✅ Documento Word final guardado en:\n{final_docx}\n')

            else:
                print('⚠️ No se encontraron textos resaltados en el PDF.')

        except Exception as e:
            print(f'❌ Error inesperado: {e}')


if __name__ == '__main__':
    main()
